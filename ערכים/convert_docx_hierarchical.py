# גרסה 2.1 - תמיכה בהיררכיה: מדור (Heading 1) -> ערכים (Heading 2)
# ללא דילוג על עמודים + טיפול משופר בציון מקורות

import json
import re
import os
import sys
from pathlib import Path

# ========== הגדרות ==========
OUTPUT_JSON = "data.json"
OUTPUT_JS = "data.js"

# הגדרות Firebase
FIREBASE_CONFIG = {
    "project_id": "",  # למשל: "my-dictionary-app"
    "database_url": "",  # למשל: "https://my-dictionary-app-default-rtdb.firebaseio.com"
    "credentials_path": ""  # נתיב לקובץ serviceAccountKey.json
}
# ============================

def select_file_gui():
    """פותח חלון לבחירת קובץ Word"""
    try:
        import tkinter as tk
        from tkinter import filedialog
        
        # יצירת חלון נסתר
        root = tk.Tk()
        root.withdraw()
        root.attributes('-topmost', True)
        
        # פתיחת דיאלוג בחירת קובץ
        file_path = filedialog.askopenfilename(
            title="בחר קובץ Word להמרה",
            filetypes=[
                ("Word Documents", "*.docx"),
                ("All files", "*.*")
            ],
            initialdir=os.path.expanduser("~")
        )
        
        root.destroy()
        
        if file_path:
            return file_path
        else:
            print("❌ לא נבחר קובץ")
            return None
            
    except ImportError:
        print("⚠️  tkinter לא זמין, נא להזין נתיב ידנית:")
        file_path = input("נתיב לקובץ Word: ").strip().strip('"')
        if os.path.exists(file_path):
            return file_path
        else:
            print(f"❌ הקובץ לא נמצא: {file_path}")
            return None

def install_dependencies():
    """בדיקה והתקנת תלויות"""
    try:
        import docx
    except ImportError:
        print("📦 מתקין python-docx...")
        os.system(f"{sys.executable} -m pip install python-docx")
        import docx
    return True

# ייבוא אחרי בדיקת תלויות
install_dependencies()
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_text_display(text):
    """ניקוי טקסט להצגה"""
    if not text:
        return ""
    return re.sub(r'\s+', ' ', text.replace('\t', ' ')).strip()

def get_heading_level(para):
    """
    מחזיר את רמת הכותרת:
    1 = מדור (Heading 1)
    2 = ערך (Heading 2)
    0 = לא כותרת
    """
    if not para.style:
        return 0
        
    style_name = para.style.name.lower()
    
    # בדיקה ישירה של שם הסגנון
    if style_name in ['heading 1', 'כותרת 1']:
        return 1
    elif style_name in ['heading 2', 'כותרת 2']:
        return 2
    elif 'heading 1' in style_name:
        return 1
    elif 'heading 2' in style_name:
        return 2
    
    # בדיקה לפי outline level ב-XML
    try:
        from docx.oxml.ns import qn
        pPr = para._element.find(qn('w:pPr'))
        if pPr is not None:
            outlineLvl = pPr.find(qn('w:outlineLvl'))
            if outlineLvl is not None:
                level = int(outlineLvl.get(qn('w:val'))) + 1
                if level in [1, 2]:
                    return level
    except:
        pass
    
    return 0

def is_source_line(para):
    """
    בדיקה האם זו שורת מקור:
    1. מתחילה ב-[ ומכילה ]
    2. מיושרת לשמאל ויזואלית (= RIGHT ב-XML של Word בעברית RTL)
    
    שורת מקור יכולה להופיע:
    - בסוף אותה פסקה של התוכן
    - בפסקה נפרדת (אפילו בעמוד הבא אחרי page break)
    """
    text = para.text.strip()
    if not (text.startswith('[') and ']' in text):
        return False
    
    # בעברית RTL, יישור ויזואלי לשמאל = RIGHT לוגית
    if para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return True
    
    # בדיקת XML ישירה
    xml = para._element.xml
    if 'w:jc' in xml and 'w:val="right"' in xml:
        return True
    
    return False

def is_brackets_line(text):
    """בדיקה האם שורה מתחילה ב-[ ומכילה ] - ללא קשר ליישור"""
    text = text.strip()
    return text.startswith('[') and ']' in text

def is_quote_paragraph(para):
    """
    בודק אם פסקה היא ציטוט מתוך מקור:
    - רוב הטקסט (מבחינת מס' תווים) הוא בגודל 10.5pt או 11.5pt
    - או בגופן Aptos (שונה מהגופן הברירת מחדל של הטקסט הראשי)
    """
    quote_chars = 0
    total_chars = 0
    for run in para.runs:
        text = run.text
        if not text.strip():
            continue
        total_chars += len(text)
        size = run.font.size.pt if run.font.size else None
        name = run.font.name
        if size in (10.5, 11.5) or name == 'Aptos':
            quote_chars += len(text)
    if total_chars == 0:
        return False
    return quote_chars / total_chars > 0.5

def para_to_html(para):
    """ממיר פסקה ל-HTML תוך שמירה על עיצוב.
    אם הפסקה כולה היא ציטוט (גופן/גודל שונים מהראשי) — עוטף ב-<span class="quote">."""
    html_content = ""

    for run in para.runs:
        text = run.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if not text:
            continue

        prefix = ""
        suffix = ""

        # בדיקת גודל גופן - אם קטן מ-11pt -> הערה קטנה
        if run.font.size:
            pt = run.font.size.pt
            if pt <= 10:
                prefix += '<span class="text-small">'
                suffix = '</span>' + suffix

        if run.bold:
            prefix += "<strong>"
            suffix = "</strong>" + suffix

        if run.underline:
            prefix += "<u>"
            suffix = "</u>" + suffix

        html_content += prefix + text + suffix

    # אם רוב הפסקה הוא בגופן/גודל ציטוט - לעטוף את הכל
    if html_content and is_quote_paragraph(para):
        html_content = f'<span class="quote">{html_content}</span>'

    return html_content

def convert_docx_to_hierarchical(filename):
    """
    ממיר מסמך DOCX למבנה היררכי: מדורים -> ערכים
    
    הלוגיקה לציון מקורות:
    - שורת מקור מזוהה ע"י: מתחילה ב-[ ומיושרת לשמאל
    - שורת מקור מחוברת לתוכן הקודם, גם אם היא בעמוד אחר (אחרי page break)
    - זה עובד כי doc.paragraphs מחזיר את כל הפסקאות ברצף, ללא קשר ל-page breaks
    """
    try:
        doc = docx.Document(filename)
    except Exception as e:
        print(f"❌ שגיאה בפתיחת הקובץ: {e}")
        return None

    # מידע על המסמך
    total_paragraphs = len(doc.paragraphs)
    print(f"\n📄 נמצאו {total_paragraphs} פסקאות במסמך")
    
    data = {
        "sections": [],
        "metadata": {
            "source_file": os.path.basename(filename),
            "version": "2.1",
            "total_paragraphs": total_paragraphs
        }
    }
    
    current_section = None
    current_term = None
    section_counter = 0
    term_counter = 0
    pending_content = ""
    source_attachments = 0  # מונה כמה מקורות חוברו לתוכן קודם

    print(f"🔄 מעבד את {os.path.basename(filename)}...")
    print(f"   (מעבד את כל הפסקאות, ללא דילוג)")

    for i, para in enumerate(doc.paragraphs):
        raw_text = clean_text_display(para.text)
        if not raw_text:
            continue
        
        html_text = para_to_html(para)
        heading_level = get_heading_level(para)

        # Heading 1 = מדור חדש
        if heading_level == 1:
            # שמירת מדור קודם
            if current_section:
                if current_term:
                    if pending_content:
                        current_term["definition"].append(pending_content)
                        pending_content = ""
                    current_section["terms"].append(current_term)
                    current_term = None
                data["sections"].append(current_section)
            
            section_counter += 1
            current_section = {
                "id": f"s_{section_counter}",
                "name": raw_text,
                "terms": []
            }
            print(f"   📁 מדור {section_counter}: {raw_text}")
        
        # Heading 2 = ערך חדש
        elif heading_level == 2:
            if current_section:
                # שמירת ערך קודם
                if current_term:
                    if pending_content:
                        current_term["definition"].append(pending_content)
                        pending_content = ""
                    current_section["terms"].append(current_term)
                
                term_counter += 1
                current_term = {
                    "id": f"t_{term_counter}",
                    "term": raw_text,
                    "definition": []
                }
            else:
                # אם אין מדור, ניצור מדור ברירת מחדל
                section_counter += 1
                current_section = {
                    "id": f"s_{section_counter}",
                    "name": "כללי",
                    "terms": []
                }
                print(f"   📁 מדור ברירת מחדל: כללי")
                
                term_counter += 1
                current_term = {
                    "id": f"t_{term_counter}",
                    "term": raw_text,
                    "definition": []
                }
        
        # תוכן רגיל (לא כותרת)
        else:
            if current_term:
                # ========================================
                # טיפול בשורת מקור (מיושרת לשמאל עם [])
                # ========================================
                if is_source_line(para):
                    source_text = raw_text
                    
                    # מחבר את המקור לתוכן הקודם
                    if pending_content:
                        # יש תוכן ממתין - מחבר אליו
                        current_term["definition"].append(pending_content + " " + source_text)
                        pending_content = ""
                        source_attachments += 1
                    elif len(current_term["definition"]) > 0:
                        # אין תוכן ממתין, אבל יש הגדרות קודמות - מחבר לאחרונה
                        # זה המקרה שבו המקור בעמוד הבא (אחרי page break)
                        current_term["definition"][-1] += " " + source_text
                        source_attachments += 1
                    else:
                        # אין כלום קודם - מוסיף כשורה בודדת (מקרה קצה)
                        current_term["definition"].append(source_text)
                
                # ========================================
                # שורה עם [] אבל לא מיושרת לשמאל - חלק מהתוכן
                # ========================================
                elif is_brackets_line(raw_text):
                    if pending_content:
                        pending_content += "<br/>" + html_text
                    else:
                        pending_content = html_text
                
                # ========================================
                # תוכן רגיל
                # ========================================
                else:
                    # אם מתחיל בסוגריים עגולים - המשך לפסקה קודמת (הערה)
                    if raw_text.startswith('(') and len(current_term["definition"]) > 0:
                        current_term["definition"][-1] += " " + html_text
                    else:
                        if pending_content:
                            html_text = pending_content + "<br/>" + html_text
                            pending_content = ""
                        current_term["definition"].append(html_text)

    # שמירת מדור וערך אחרונים
    if current_section:
        if current_term:
            if pending_content:
                current_term["definition"].append(pending_content)
            current_section["terms"].append(current_term)
        data["sections"].append(current_section)

    data["metadata"]["source_attachments"] = source_attachments
    print(f"\n   ✓ עובדו {total_paragraphs} פסקאות")
    print(f"   ✓ חוברו {source_attachments} ציוני מקור לתוכן הקודם")
    
    return data

def save_to_json(data, filename):
    """שמירה לקובץ JSON"""
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"💾 נשמר: {filename}")

def save_to_js(data, filename):
    """שמירה לקובץ JS"""
    with open(filename, "w", encoding="utf-8") as f:
        f.write("const rawData = " + json.dumps(data, ensure_ascii=False, indent=2) + ";")
    print(f"💾 נשמר: {filename}")

def upload_to_firebase(data):
    """העלאה ל-Firebase Realtime Database"""
    if not FIREBASE_CONFIG["project_id"] or not FIREBASE_CONFIG["credentials_path"]:
        print("\n⚠️  Firebase לא מוגדר - מדלג על העלאה")
        print("   (ראה הוראות בקובץ FIREBASE_SETUP.md)")
        return False
    
    if not os.path.exists(FIREBASE_CONFIG["credentials_path"]):
        print(f"\n❌ קובץ credentials לא נמצא: {FIREBASE_CONFIG['credentials_path']}")
        return False
    
    try:
        import firebase_admin
        from firebase_admin import credentials, db
        
        print("\n🔥 מתחבר ל-Firebase...")
        
        # אתחול Firebase (רק אם לא אותחל כבר)
        if not firebase_admin._apps:
            cred = credentials.Certificate(FIREBASE_CONFIG["credentials_path"])
            firebase_admin.initialize_app(cred, {
                'databaseURL': FIREBASE_CONFIG["database_url"]
            })
        
        # העלאת הנתונים
        ref = db.reference('/dictionary')
        ref.set(data)
        
        print(f"✅ הועלה בהצלחה ל-Firebase!")
        print(f"   URL: {FIREBASE_CONFIG['database_url']}/dictionary")
        return True
        
    except ImportError:
        print("\n⚠️  חסרה ספריית firebase-admin")
        print("   להתקנה: pip install firebase-admin")
        return False
    except Exception as e:
        print(f"\n❌ שגיאה בהעלאה ל-Firebase: {e}")
        return False

def print_summary(data):
    """הדפסת סיכום"""
    total_sections = len(data["sections"])
    total_terms = sum(len(s["terms"]) for s in data["sections"])
    
    print("\n" + "="*60)
    print("📊 סיכום:")
    print(f"   📁 מדורים: {total_sections}")
    print(f"   📝 ערכים: {total_terms}")
    print("="*60)
    
    for section in data["sections"]:
        print(f"\n📁 {section['name']} ({len(section['terms'])} ערכים)")
        for term in section["terms"][:3]:
            print(f"   • {term['term']}")
        if len(section["terms"]) > 3:
            print(f"   ... ועוד {len(section['terms']) - 3} ערכים")

def main():
    print("="*60)
    print("  📚 ממיר מילון טללי חיים - גרסה 2.1")
    print("="*60)
    
    # בחירת קובץ
    print("\n📂 בחר קובץ Word להמרה...")
    input_file = select_file_gui()
    
    if not input_file:
        return
    
    print(f"\n✓ נבחר: {input_file}")
    
    # קביעת תיקיית פלט (אותה תיקייה כמו הקובץ)
    output_dir = os.path.dirname(input_file)
    output_json = os.path.join(output_dir, OUTPUT_JSON)
    output_js = os.path.join(output_dir, OUTPUT_JS)
    
    # המרה
    data = convert_docx_to_hierarchical(input_file)
    
    if data is None:
        return
    
    # שמירה
    print("\n💾 שומר קבצים...")
    save_to_json(data, output_json)
    save_to_js(data, output_js)
    
    # העלאה ל-Firebase
    upload_to_firebase(data)
    
    # סיכום
    print_summary(data)
    
    print(f"\n✅ הסתיים בהצלחה!")
    print(f"   קבצי הפלט נשמרו ב: {output_dir}")
    
    # המתנה לפני סגירה
    input("\nלחץ Enter לסגירה...")

if __name__ == "__main__":
    main()
