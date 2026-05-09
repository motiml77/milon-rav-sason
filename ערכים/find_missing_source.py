# -*- coding: utf-8 -*-
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

folder = os.path.dirname(__file__)
docx_files = [f for f in os.listdir(folder) if f.endswith('.docx')]
docx_path = os.path.join(folder, docx_files[0])
doc = docx.Document(docx_path)

target = 'כוחם של ישראל, בחינת נשמות דדכורא'
for pi, para in enumerate(doc.paragraphs):
    if target in para.text:
        print(f'\n=== Paragraph {pi}: {para.text[:200]!r} ===')
        print(f'  alignment: {para.alignment}')
        # Show next 5 paragraphs
        for offset in range(1, 6):
            if pi + offset < len(doc.paragraphs):
                np = doc.paragraphs[pi + offset]
                t = np.text
                if not t.strip():
                    print(f'  [{pi+offset}] EMPTY')
                    continue
                # Check style and alignment
                style = np.style.name if np.style else 'no-style'
                align = np.alignment
                # Check XML for jc=right
                xml = np._element.xml
                jc_right = 'w:val="right"' in xml
                print(f'  [{pi+offset}] style={style} align={align} jc_right={jc_right}')
                print(f'      text: {t[:150]!r}')
        break
