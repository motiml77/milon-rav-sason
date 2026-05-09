# -*- coding: utf-8 -*-
"""חיפוש ציטוט ספציפי במסמך וניתוח הגופנים שלו"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import docx

folder = os.path.dirname(__file__)
docx_files = [f for f in os.listdir(folder) if f.endswith('.docx')]
docx_path = os.path.join(folder, docx_files[0])
doc = docx.Document(docx_path)

# Search for the quote
search_terms = ['ראוי לדעת למה קראו', 'מפורש ביותר כתב בספר', 'תפילה למשה']

for pi, para in enumerate(doc.paragraphs):
    text = para.text
    for term in search_terms:
        if term in text:
            print(f'\n=== Paragraph {pi} ===')
            print(f'Full text: {text[:300]!r}')
            print('Run breakdown:')
            for ri, run in enumerate(para.runs):
                if not run.text.strip():
                    continue
                size = run.font.size.pt if run.font.size else None
                name = run.font.name or 'default'
                italic = run.italic
                bold = run.bold
                print(f'  [{ri}] ({name}, {size}pt, bold={bold}, italic={italic}): {run.text[:80]!r}')
            print()
            # also show next 3 paragraphs
            for offset in range(1, 4):
                if pi + offset < len(doc.paragraphs):
                    np = doc.paragraphs[pi + offset]
                    nt = np.text[:100]
                    if not nt.strip():
                        continue
                    print(f'--- Next para {pi+offset}: {nt!r}')
                    for run in np.runs:
                        if not run.text.strip():
                            continue
                        size = run.font.size.pt if run.font.size else None
                        name = run.font.name or 'default'
                        print(f'    ({name}, {size}pt): {run.text[:60]!r}')
            break
