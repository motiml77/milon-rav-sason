# -*- coding: utf-8 -*-
"""מראה פסקאות שלמות בהן יש ערבוב גדלים, להבין את ההקשר של 10.5/11.5 ו-Aptos"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
import docx

folder = os.path.dirname(__file__)
docx_files = [f for f in os.listdir(folder) if f.endswith('.docx')]
docx_path = os.path.join(folder, docx_files[0])
doc = docx.Document(docx_path)

# Find paragraphs containing 10.5pt runs and show structure
print('=== Paragraphs with 10.5pt runs ===\n')
shown = 0
for pi, para in enumerate(doc.paragraphs):
    has_105 = any(r.font.size and r.font.size.pt == 10.5 for r in para.runs if r.text.strip())
    if has_105 and shown < 4:
        print(f'--- Paragraph {pi} ---')
        for r in para.runs:
            if not r.text.strip():
                continue
            size = r.font.size.pt if r.font.size else None
            name = r.font.name or 'default'
            marker = '★' if size == 10.5 else ('●' if size and size < 10 else '·')
            txt = r.text[:90]
            print(f'  {marker} ({name}, {size}): {txt!r}')
        print()
        shown += 1

print('\n=== Paragraphs with 11.5pt runs ===\n')
shown = 0
for pi, para in enumerate(doc.paragraphs):
    has_115 = any(r.font.size and r.font.size.pt == 11.5 for r in para.runs if r.text.strip())
    if has_115 and shown < 4:
        print(f'--- Paragraph {pi} ---')
        for r in para.runs:
            if not r.text.strip():
                continue
            size = r.font.size.pt if r.font.size else None
            name = r.font.name or 'default'
            marker = '★' if size == 11.5 else ('●' if size and size < 10 else '·')
            txt = r.text[:90]
            print(f'  {marker} ({name}, {size}): {txt!r}')
        print()
        shown += 1

print('\n=== Paragraphs with Aptos runs (non-source) ===\n')
shown = 0
for pi, para in enumerate(doc.paragraphs):
    aptos_runs = [r for r in para.runs if r.font.name == 'Aptos' and r.text.strip()]
    if aptos_runs and shown < 5:
        # Skip if all aptos runs are source brackets
        all_brackets = all('[' in r.text or ']' in r.text or 'ראה' in r.text for r in aptos_runs)
        if all_brackets:
            continue
        print(f'--- Paragraph {pi} ---')
        for r in para.runs:
            if not r.text.strip():
                continue
            size = r.font.size.pt if r.font.size else None
            name = r.font.name or 'default'
            marker = '★' if name == 'Aptos' else '·'
            txt = r.text[:90]
            print(f'  {marker} ({name}, {size}): {txt!r}')
        print()
        shown += 1
